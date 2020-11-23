from xlrd import open_workbook
from docx import Document
from os import path, mkdir, remove
from datetime import datetime
from locale import setlocale, LC_ALL, format_string
from num2words import num2words
import sys
import re
from PyPDF2 import PdfFileReader, PdfFileWriter
from glob import glob
from time import sleep
import comtypes.client
from fpdf import FPDF 
from tkinter import Tk, filedialog, Label, Grid, Button, Entry, OptionMenu, StringVar


class PeticaoEmLote:
    modelo = ''
    planilha = ''
    ac = ''
    genero = ''
    resultado = ''


def create_modified_docx(template, planilha, ac, genero, output, window):
    if genero.lower() == 'h':
        cargo = 'Procurador'
    elif genero.lower() == 'm':
        cargo = 'Procuradora'
    else:
        print('Favor informar H/M para o genero')
        sys.exit(2)
    wb = open_workbook(planilha)
    sheet = wb.sheet_by_index(0)

    for line in range(1, sheet.nrows - 1):
    # for line in range(1, 2):
        document = Document(template)
        uf = sheet.cell_value(line, 0)
        cnpj = sheet.cell_value(line, 1)
        filial = "%s.%s.%s/%s-%s" % (cnpj[0:2], cnpj[2:5], cnpj[5:8], cnpj[8:12], cnpj[12:14])
        ie = sheet.cell_value(line, 2)
        local = sheet.cell_value(line, 3)
        local_caps = local.upper()
        endereco = sheet.cell_value(line, 4)
        mes = sheet.cell_value(line, 5)
        ano = sheet.cell_value(line, 6)
        setlocale(LC_ALL, 'pt-BR')
        date_object = datetime(month=int(mes), year=int(ano), day=1)
        periodo = date_object.strftime('%B de %Y')
        path_periodo = date_object.strftime('%m%Y')
        valor = sheet.cell_value(line, 7)
        correspondente = sheet.cell_value(line, 8)

        docx_replace(document, {
                                '{{uf_extenso}}': uf_por_extenso(uf).upper(),
                                '{{uf}}': uf,
                                '{{filial}}': filial,
                                '{{ie}}': ie,
                                '{{local}}': local,
                                '{{local_caps}}': local_caps,
                                '{{endereco}}': endereco,
                                '{{mes}}': mes,
                                '{{ano}}': ano,
                                '{{periodo}}': periodo,
                                '{{valor}}': format_string('R$ %.2f', valor, grouping=True),
                                '{{valor_extenso}}': num2words(valor, lang='pt_BR', to='currency'),
                                '{{data_atual}}': datetime.now().strftime('%d de %B de %Y'),
                                '{{ac}}': ac,
                                '{{cargo}}': cargo,
                                '{{correspondente}}': correspondente,
                                })
        if not path.isdir(output):
            raise "O diretório de saída não existe"
        result_path = path.join(output, cnpj, path_periodo)
        result_name = f'001-peticao'
        document.save(f'{path.join(result_path,result_name)}.docx')
        convert_to_pdf(f'{path.join(result_path,result_name)}.docx', f'{path.join(result_path,result_name)}.pdf')
        if path.isfile(f'{path.join(result_path,result_name)}_pronta.pdf'):
            remove(f'{path.join(result_path,result_name)}_pronta.pdf')
        pdf_cat(get_pdfs(result_path), output_stream=open(f'{path.join(result_path,result_name)}_pronta.pdf', 'w+b'))
    Label(window, text=f"Todos os períodos finalizados").grid(row=line+5, column=0)


def uf_por_extenso(uf):
    uf_object = {
        'AC': 'Acre',
        'AL': 'Alagoas',
        'AP': 'Amapá',
        'AM': 'Amazonas',
        'BA': 'Bahia',
        'CE': 'Ceará',
        'DF': 'Distrito Federal',
        'ES': 'Espírito Santo',
        'GO': 'Goiás',
        'MA': 'Maranhão',
        'MT': 'Mato Grosso',
        'MS': 'Mato Grosso do Sul',
        'MG': 'Minas Gerais',
        'PA': 'Pará',
        'PB': 'Paraíba',
        'PR': 'Paraná',
        'PE': 'Pernambuco',
        'PI': 'Piauí',
        'RJ': 'Rio de Janeiro',
        'RN': 'Rio Grande do Norte',
        'RS': 'Rio Grande do Sul',
        'RO': 'Rondônia',
        'RR': 'Roraima',
        'SC': 'Santa Catarina',
        'SP': 'São Paulo',
        'SE': 'Sergipe',
        'TO': 'Tocantins'
    }
    return uf_object.get(uf)


def docx_replace(doc, data):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        for key, val in data.items():
            key_name = key # I'm using placeholders in the form ${PlaceholderName}
            if key_name in p.text:
                inline = p.runs


                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):
                    if re.search(key_name, inline[i].text):
                    #  if key_name == inline[i].text:
                        inline[i].text = re.sub(key_name, str(val), inline[i].text)
                        # found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        # text = inline[i].text.replace(key_name, str(val))
                        # inline[i].text = text
                        replace_done = True
                        found_all = True
                        break
                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break
                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                                # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break

                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            try:
                                if inline[i].text[text_index] == key_name[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            except:
                                continue
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break

                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text


def get_pdfs(list_to_add):
    all_files = [path.join(list_to_add,'001-peticao.pdf')]
    pdfs = sorted(glob("./pdf/*.pdf"))
    all_files += pdfs
    folder_glob = glob(f'{list_to_add}/*') 
    folder_glob_reduzido = [s for s in folder_glob if not s.__contains__('001-peticao')]
    all_files += folder_glob_reduzido
    inicio = '01-Minuta'
    inicio_glob = glob(path.join(list_to_add,inicio+'*'))
    if inicio_glob:
        inicio_glob_reduzido = [s for s in all_files if not s.__contains__(inicio)]
        all_files = inicio_glob + inicio_glob_reduzido
    return all_files


def pdf_cat(input_files, output_stream):
    input_streams = []
    try:
        # First open all the files, then produce the output file, and
        # finally close the input files. This is necessary because
        # the data isn't read from the input files until the write
        # operation. Thanks to
        # https://stackoverflow.com/questions/6773631/problem-with-closing-python-pypdf-writing-getting-a-valueerror-i-o-operation/6773733#6773733
        for input_file in input_files:
            filename, file_extension = path.splitext(input_file)
            if file_extension.lower() == '.pdf':
                input_streams.append(open(input_file, 'rb'))
            elif file_extension.lower() == '.xlsx':
                if path.isfile(filename+'.pdf'):
                    remove(filename+'.pdf')
                convert_xlsx_to_pdf(input_file, filename+'.pdf')
                input_streams.append(open(filename+'.pdf', 'rb'))
            elif file_extension.lower() == '.csv':
                if path.isfile(filename+'.pdf'):
                    remove(filename+'.pdf')
                txt_to_pdf(input_file, filename+'.csv.pdf')
                input_streams.append(open(filename+'.csv.pdf', 'rb'))
            elif file_extension.lower() == '.txt':
                if path.isfile(filename+'.pdf'):
                    remove(filename+'.pdf')
                txt_to_pdf(input_file, filename+'.txt.pdf')
                input_streams.append(open(filename+'.txt.pdf', 'rb'))
        writer = PdfFileWriter()
        for reader in map(PdfFileReader, input_streams):
            for n in range(reader.getNumPages()):
                writer.addPage(reader.getPage(n))
        writer.write(output_stream)
    except:
        raise
    finally:
        for f in input_streams:
            f.close()


def convert_to_pdf(input_file, output_file):
    try:
        wdFormatPDF = 17

        in_file = path.abspath(input_file)
        out_file = path.abspath(output_file)
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
    except:
        raise

def txt_to_pdf(input_file, output_file):
    # save FPDF() class into a  
    # variable pdf 
    pdf = FPDF('L', 'mm','A4') 
    
    # Add a page 
    pdf.add_page() 
    
    # set style and size of font  
    # that you want in the pdf 
    pdf.set_font("Arial", size = 6) 
    
    # open the text file in read mode 
    with open(input_file, "r", newline='', encoding='iso-8859-1') as f:
        # insert the texts in pdf 
        for x in f: 
            pdf.multi_cell(220,6, border=0, txt = x)

    
    # save the pdf with name .pdf 
    pdf.output(output_file)    

def convert_xlsx_to_pdf(input_file, output_file):
    app = comtypes.client.CreateObject('Excel.Application')
    app.Visible = False

    infile = path.join(path.abspath(input_file))
    outfile = path.join(path.abspath(output_file))

    doc = app.Workbooks.Open(infile)
    doc.ExportAsFixedFormat(0, outfile, 1, 0)
    
    doc.Close()

    app.Quit()

def openfilename(peticao_em_lote, atribute, extensions):
    setattr(peticao_em_lote, atribute,filedialog.askopenfilename(filetypes=extensions))

def opendirectory(peticao_em_lote):
    peticao_em_lote.resultado = filedialog.askdirectory()

def execute(peticao_em_lote, window):
    try:
        Label(window, text="Criando as petições. Aguarde.").grid(row=6, column=0)
        create_modified_docx(peticao_em_lote.modelo, peticao_em_lote.planilha, peticao_em_lote.ac.get(), peticao_em_lote.genero.get(), peticao_em_lote.resultado, window)
        Label(window, text="Todas as petições criadas com sucesso.").grid(row=7, column=0)
    except Exception as e:
        print(e)
        Label(window, text=f"Ocorreu um erro ao criar as petições:").grid(row=6)
        Label(window, text=e, width=25, wraplength=150).grid(row=7)
        pass
    print('Finalizado')

if __name__ == '__main__':
    peticao_em_lote = PeticaoEmLote()
    root = Tk()
    root.title('Criar petições em lote')
    root.geometry('500x500')

    label_modelo = Label(root, text="Selecione o modelo")
    label_modelo.grid(row=0,column=0)
    modelo = Button(root, text="Selecionar", command=lambda: openfilename(peticao_em_lote, 'modelo', [('Word', '*.docx')]), width=20)
    modelo.grid(row=0,column=1)

    label_planilha = Label(root, text="Selecione a planilha")
    label_planilha.grid(row=1,column=0)
    planilha = Button(root, text="Selecionar", command=lambda: openfilename(peticao_em_lote, 'planilha', [('Excel', '*.xlsx')]), width=20)
    planilha.grid(row=1,column=1)

    label_ac = Label(root, text="Digite o nome do destinatário:")
    label_ac.grid(row=2,column=0)
    peticao_em_lote.ac = Entry(root, width=25)
    peticao_em_lote.ac.grid(row=2,column=1)

    label_genero = Label(root, text="Selecione o genero do parceiro")
    label_genero.grid(row=3,column=0)
    peticao_em_lote.genero = StringVar()
    peticao_em_lote.genero.set('H')
    genero = OptionMenu(root, peticao_em_lote.genero, 'H', 'M')
    genero.grid(row=3,column=1)

    label_resultado = Label(root, text="Selecione o diretorio de destino")
    label_resultado.grid(row=4,column=0)
    peticao_em_lote.resultado = Button(root, text="Selecionar", command=lambda: opendirectory(peticao_em_lote), width=20)
    peticao_em_lote.resultado.grid(row=4,column=1)

    executar = Button(root, text="Executar", command=lambda: execute(peticao_em_lote, root), width=20)
    executar.grid(row=5, column=0)
    root.mainloop()
